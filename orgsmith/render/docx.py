""".docx renderer: python-docx with house style, letterhead, PAGE-field
footer, sigblocks, core properties, and the synthetic marker."""

from __future__ import annotations

import io
from datetime import datetime, timezone

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from ..fabric.engagements import render_date
from ..schemas import DocIR, ManifestEntry
from .provenance import add_docx_marker
from .styles import StylePack


def _page_number_field(paragraph) -> None:
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), " PAGE ")
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "1"
    run.append(text)
    fld.append(run)
    paragraph._p.append(fld)


def _sig_lines(doc, person, when_text: str) -> None:
    doc.add_paragraph("")
    doc.add_paragraph("_____________________________")
    name_p = doc.add_paragraph()
    name_p.add_run(person["name"]).bold = True
    doc.add_paragraph(person["title"])
    doc.add_paragraph(f"Date: {when_text}")


def render_docx(
    docir: DocIR,
    entry: ManifestEntry,
    style: StylePack,
    author_name: str,
    people: dict[str, dict],
) -> bytes:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = style.font_family
    normal.font.size = Pt(11)

    section = doc.sections[0]
    section.different_first_page_header_footer = True
    head = section.first_page_header.paragraphs[0]
    head.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = head.add_run(style.letterhead_lines[0])
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor.from_string(style.accent_hex)
    for line in style.letterhead_lines[1:]:
        p = section.first_page_header.add_paragraph(line)
        p.runs[0].font.size = Pt(9)

    later = section.header.paragraphs[0]
    later_run = later.add_run(style.letterhead_lines[0])
    later_run.font.size = Pt(9)

    for target in (section.footer, section.first_page_footer):
        fp = target.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.add_run("Page ").font.size = Pt(9)
        _page_number_field(fp)

    when_text = render_date(entry.date)
    for block in docir.blocks:
        if block.kind == "heading":
            doc.add_heading(block.text, level=min(max(block.level, 1), 4))
        elif block.kind == "paragraph":
            doc.add_paragraph(block.text)
        elif block.kind == "list":
            for item in block.items:
                doc.add_paragraph(item, style="List Bullet")
        elif block.kind == "table":
            cols = max(len(block.header), max((len(r) for r in block.rows), default=1))
            table = doc.add_table(rows=0, cols=cols)
            table.style = "Table Grid"
            if block.header:
                cells = table.add_row().cells
                for i, text in enumerate(block.header):
                    cells[i].paragraphs[0].add_run(text).bold = True
            for row in block.rows:
                cells = table.add_row().cells
                for i, text in enumerate(row):
                    cells[i].text = text
        elif block.kind == "sigblock":
            for signer in block.signers:
                _sig_lines(doc, people[signer], when_text)

    props = doc.core_properties
    props.title = entry.title
    props.author = author_name
    props.last_modified_by = author_name
    stamp = datetime(
        entry.date.year, entry.date.month, entry.date.day, 9, 0, 0,
        tzinfo=timezone.utc,
    )
    props.created = stamp
    props.modified = stamp

    buf = io.BytesIO()
    doc.save(buf)
    return add_docx_marker(buf.getvalue())
