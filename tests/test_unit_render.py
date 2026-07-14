"""Unit tier: renderers ground documents in the ledger, verified by readers."""

import re

import pytest

from orgsmith.artifacts import (
    load_engagements,
    load_finance,
    load_manifest,
)
from orgsmith.assemble import run_assemble
from orgsmith.render import run_render
from orgsmith.render.provenance import (
    docx_has_marker,
    pdf_has_marker,
    xlsx_has_marker,
)
from orgsmith.render.resolve import FactResolutionError, resolve_docir
from orgsmith.schemas import Block, DocIR

from conftest import build_pure_stages, run_authoring, run_enrichment

pytestmark = pytest.mark.unit


@pytest.fixture(scope="module")
def org(tmp_path_factory):
    paths = build_pure_stages(tmp_path_factory.mktemp("rendered-org"))
    run_enrichment(paths)
    run_authoring(paths)
    assert run_render(paths) == 0
    assert run_assemble(paths) == 0
    return paths


def _docx_text(path) -> str:
    import docx

    doc = docx.Document(str(path))
    chunks = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            chunks.extend(cell.text for cell in row.cells)
    return "\n".join(chunks)


def _pdf_text(path) -> str:
    from pypdf import PdfReader

    return "\n".join(page.extract_text() or "" for page in PdfReader(str(path)).pages)


def _norm(text: str) -> str:
    return re.sub(r"\s+", " ", text)


def test_every_manifest_doc_rendered_and_listed(org):
    manifest = load_manifest(org)
    toc = org.toc_md.read_text()
    for entry in manifest:
        assert (org.share_dir / entry.path).exists(), entry.path
        assert entry.path.rsplit("/", 1)[-1] in toc


def test_fact_echo_docx_and_pdf(org):
    manifest = load_manifest(org)
    facts = load_engagements(org).fact_index()
    checked = 0
    for entry in manifest:
        if not entry.facts_refs or entry.format == "xlsx":
            continue
        path = org.share_dir / entry.path
        text = _norm(_docx_text(path) if entry.format == "docx" else _pdf_text(path))
        for ref in entry.facts_refs:
            assert facts[ref].rendered in text, f"{entry.path}: {ref}"
            checked += 1
    assert checked > 10


def test_provenance_markers_all_formats(org):
    manifest = load_manifest(org)
    for entry in manifest:
        path = org.share_dir / entry.path
        has = {
            "docx": docx_has_marker,
            "pdf": pdf_has_marker,
            "xlsx": xlsx_has_marker,
        }[entry.format]
        assert has(path), entry.path


def _eval_formula(formula: str, cached) -> int:
    body = formula.lstrip("=")
    m = re.fullmatch(r"SUM\(([A-Z]+\d+):([A-Z]+\d+)\)", body)
    if m:
        cells = list(cached[m.group(1) : m.group(2)])
        return sum(c.value or 0 for row in cells for c in row)
    m = re.fullmatch(r"([A-Z]+\d+)-([A-Z]+\d+)", body)
    if m:
        return (cached[m.group(1)].value or 0) - (cached[m.group(2)].value or 0)
    raise AssertionError(f"formula outside restricted vocabulary: {formula}")


def test_xlsx_formulas_recompute_to_cached_ledger_values(org):
    from openpyxl import load_workbook

    manifest = load_manifest(org)
    finance = load_finance(org)
    workbooks = [e for e in manifest if e.format == "xlsx"]
    assert len(workbooks) == 2
    for entry in workbooks:
        path = org.share_dir / entry.path
        formulas = load_workbook(path, data_only=False)["Summary"]
        cached = load_workbook(path, data_only=True)["Summary"]

        fy = next(y for y in finance.years if y.year == int(
            entry.render_params["year"]))
        formula_cells = 0
        for row in formulas.iter_rows():
            for cell in row:
                if isinstance(cell.value, str) and cell.value.startswith("="):
                    got = cached[cell.coordinate].value
                    assert got is not None, f"no cached value at {cell.coordinate}"
                    assert _eval_formula(cell.value, cached) == got
                    formula_cells += 1
        assert formula_cells == 3  # revenue SUM, expenses SUM, net income
        # Cached revenue total ties to the ledger.
        assert cached["F4"].value == fy.revenue
        assert [cached[f"{c}4"].value for c in "BCDE"] == fy.quarters


def test_unresolved_placeholder_fails_loudly(org):
    facts = load_engagements(org).fact_index()
    doc = DocIR(
        doc_id="d:0001",
        blocks=[Block(kind="paragraph", text="fee is {{fact:not.a.fact}}")],
    )
    with pytest.raises(FactResolutionError):
        resolve_docir(doc, facts)
    doc = DocIR(
        doc_id="d:0001",
        blocks=[Block(kind="paragraph", text="malformed {{fact:}} placeholder")],
    )
    with pytest.raises(FactResolutionError):
        resolve_docir(doc, facts)


def test_render_is_incremental(org):
    manifest = load_manifest(org)
    targets = [org.share_dir / e.path for e in manifest]
    before = {p: p.stat().st_mtime_ns for p in targets}
    assert run_render(org) == 0
    after = {p: p.stat().st_mtime_ns for p in targets}
    assert before == after
